#!/usr/bin/env python3
"""
{{FIRM_NAME}} Pleading Finalizer
=======================
Takes a drafted .docx and produces the filing-ready Word document plus a PDF.

Why python-docx and not the Node `docx` library used elsewhere in this skill
library: Node `docx` is generation-only — it cannot open an existing file. This
script must READ a draft and normalize it in place, so python-docx is the correct
tool. Generation-from-scratch skills should keep using the Node builders.

What it enforces ({{FIRM_NAME}} standard, current as of 07.31.26):
  portrait · Times New Roman 12 · all black · single spacing · 1" margins
  caption: court centered / case + division right / parties left / rule line + slash
  title: centered, bold, ALL CAPS
  numbering hierarchy: I. → 1. → a. → i.

What it strips:
  highlights · non-black text · tracked changes · comments · drafter notes ·
  AI instruction blocks · clause-bank labels · option markers

What it REPORTS but never silently removes:
  bracketed placeholders — a human decides whether a blank is intentional

Usage:
  python3 finalize_pleading.py IN.docx OUT.docx [--title "..."] [--no-pdf]
Exit codes:
  0 clean · 2 finalized with placeholders remaining (review required)
"""

import sys, os, re, subprocess, argparse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

FONT = "Times New Roman"
SIZE = Pt(12)
LINE_SPACING = 1.0          # {{FIRM_NAME}} standard: single spacing, 0pt before/after
BLACK = RGBColor(0, 0, 0)

# Lines that must never survive into a filed document.
ARTIFACT_PATTERNS = [
    r"^\s*AI INSTRUCTION",
    r"^\s*\[?DRAFTER NOTE",
    r"^\s*\[?INTERNAL NOTE",
    r"^\s*\[?NOTE TO (DRAFTER|ATTORNEY|SELF)",
    r"^\s*\[?CLAUSE\s*(BANK|OPTION|MODULE)",
    r"^\s*MODULE\s+[A-Z0-9]+\s*[-—:]",
    r"^\s*\[?OPTION\s+[AB]\b",
    r"^\s*<!--",
    r"^\s*//\s",
    r"^\s*TODO\b",
]
PLACEHOLDER_RE = re.compile(r"\[[^\]]{1,120}\]|_{3,}|＿{3,}")


# ----------------------------------------------------------------- utilities
def delete_paragraph(p):
    el = p._element
    el.getparent().remove(el)
    p._p = p._element = None


def strip_revisions_and_comments(doc):
    """Accept tracked insertions, drop deletions, remove comment anchors."""
    body = doc.element.body
    removed = {"ins": 0, "del": 0, "comments": 0}

    for ins in body.findall(f".//{qn('w:ins')}"):
        parent = ins.getparent()
        idx = list(parent).index(ins)
        for child in list(ins):                      # unwrap: keep the text
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)
        removed["ins"] += 1

    for dele in body.findall(f".//{qn('w:del')}"):   # drop deleted text entirely
        dele.getparent().remove(dele)
        removed["del"] += 1

    for tag in ("w:commentRangeStart", "w:commentRangeEnd", "w:commentReference"):
        for el in body.findall(f".//{qn(tag)}"):
            el.getparent().remove(el)
            removed["comments"] += 1

    return removed


def normalize_run(run):
    """Force font, size, and true black. Clear every highlight."""
    run.font.name = FONT
    run.font.size = SIZE
    run.font.color.rgb = BLACK
    run.font.highlight_color = None
    rpr = run._element.get_or_add_rPr()
    for tag in ("w:highlight", "w:shd"):
        for el in rpr.findall(qn(tag)):
            rpr.remove(el)
    rfonts = rpr.find(qn("w:rFonts"))               # East Asian / complex script
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def normalize_paragraph(p):
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = False
    # shading on the paragraph itself
    ppr = p._element.get_or_add_pPr()
    for el in ppr.findall(qn("w:shd")):
        ppr.remove(el)
    for r in p.runs:
        normalize_run(r)


def set_portrait_and_margins(doc):
    for s in doc.sections:
        if s.orientation != WD_ORIENT.PORTRAIT or s.page_width > s.page_height:
            s.orientation = WD_ORIENT.PORTRAIT
            s.page_width, s.page_height = min(s.page_width, s.page_height), \
                                          max(s.page_width, s.page_height)
        s.left_margin = s.right_margin = Inches(1)
        s.top_margin = s.bottom_margin = Inches(1)


def is_artifact(text):
    return any(re.search(pat, text, re.I) for pat in ARTIFACT_PATTERNS)


# ------------------------------------------------------------------- caption
def apply_caption_format(doc, limit=25):
    """
    {{FIRM_NAME}} caption layout:
        IN THE CIRCUIT COURT ...            centered
        CASE NO.: / DIVISION:               right justified
        PETITIONER, / Petitioner,           left justified
        and / RESPONDENT, / Respondent.     left justified
        ______________________/            rule line ending in slash
        TITLE OF DOCUMENT                   centered, bold, ALL CAPS
    Only the first `limit` paragraphs are treated as caption territory.
    """
    applied = []
    for i, p in enumerate(doc.paragraphs[:limit]):
        t = p.text.strip()
        if not t:
            continue
        up = t.upper()
        if up.startswith("IN THE CIRCUIT COURT") or up.startswith("IN THE COUNTY COURT") \
           or "JUDICIAL CIRCUIT" in up or re.match(r"^IN AND FOR .*COUNTY", up):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            applied.append(("court-centered", t[:48]))
        elif re.match(r"^(CASE\s*(NO|NUMBER)|DIVISION|DIV\b|JUDGE)\b", up):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            applied.append(("case-right", t[:48]))
        elif re.match(r"^_{5,}\s*/?\s*$", t) or t.endswith("/") and set(t) <= set("_/ "):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if not t.rstrip().endswith("/"):        # slash is mandatory
                for r in p.runs:
                    r.text = ""
                p.runs[0].text = "_" * 40 + "/" if p.runs else None
                normalize_run(p.runs[0])
            applied.append(("rule-line", "____/"))
        elif re.match(r"^(Petitioner|Respondent|Former Wife|Former Husband|Plaintiff|Defendant)[,.]?$", t, re.I) \
             or up in ("AND",):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            applied.append(("party-role-left", t[:48]))
    return applied


def format_title(doc, title_text=None, limit=30):
    """Centre and bold the document title — the first ALL-CAPS line after the rule."""
    seen_rule = False
    for p in doc.paragraphs[:limit]:
        t = p.text.strip()
        if not t:
            continue
        if re.match(r"^_{5,}", t):
            seen_rule = True
            continue
        if seen_rule:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.text = r.text.upper()
                normalize_run(r)
            return t
    if title_text:                                   # nothing found — insert one
        return None
    return None


# ---------------------------------------------------------------------- main
def finalize(src, dst, title=None, make_pdf=True):
    doc = Document(src)
    report = {"artifacts_removed": [], "placeholders": [], "caption": [],
              "revisions": {}, "title": None, "paragraphs": 0}

    report["revisions"] = strip_revisions_and_comments(doc)
    set_portrait_and_margins(doc)

    for p in list(doc.paragraphs):                   # artifact lines
        if p.text.strip() and is_artifact(p.text):
            report["artifacts_removed"].append(p.text.strip()[:90])
            delete_paragraph(p)

    for p in doc.paragraphs:                         # normalize everything
        normalize_paragraph(p)
        report["paragraphs"] += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    normalize_paragraph(p)

    report["caption"] = apply_caption_format(doc)
    report["title"] = format_title(doc, title)

    for p in doc.paragraphs:                         # report placeholders, don't touch
        for m in PLACEHOLDER_RE.findall(p.text):
            if m.strip("_ ") == "" and len(m) > 20:  # signature/rule lines are fine
                continue
            report["placeholders"].append(m[:80])

    doc.save(dst)

    pdf = None
    if make_pdf:
        outdir = os.path.dirname(os.path.abspath(dst)) or "."
        res = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, dst],
            capture_output=True, text=True, timeout=180)
        cand = os.path.join(outdir, os.path.splitext(os.path.basename(dst))[0] + ".pdf")
        pdf = cand if os.path.exists(cand) else None
        if pdf is None:
            report["pdf_error"] = (res.stderr or res.stdout)[:300]
    return report, pdf


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("src"); ap.add_argument("dst")
    ap.add_argument("--title", default=None)
    ap.add_argument("--no-pdf", action="store_true")
    a = ap.parse_args()

    rep, pdf = finalize(a.src, a.dst, a.title, make_pdf=not a.no_pdf)

    print("=" * 62)
    print(" {{FIRM_NAME}} PLEADING FINALIZER")
    print("=" * 62)
    print(f" Source : {a.src}")
    print(f" Word   : {a.dst}")
    print(f" PDF    : {pdf or 'NOT GENERATED — ' + rep.get('pdf_error','')}")
    print(f"\n Enforced: portrait · {FONT} 12 · all black · {LINE_SPACING} spacing · 1\" margins")
    print(f" Paragraphs normalized: {rep['paragraphs']}")

    rv = rep["revisions"]
    print(f"\n Tracked changes accepted: {rv['ins']} | deletions dropped: {rv['del']} "
          f"| comment anchors removed: {rv['comments']}")

    print(f"\n Artifact lines removed: {len(rep['artifacts_removed'])}")
    for x in rep["artifacts_removed"]:
        print(f"   - {x}")

    print(f"\n Caption elements formatted: {len(rep['caption'])}")
    for kind, txt in rep["caption"]:
        print(f"   - {kind:18} {txt}")
    print(f" Title: {rep['title'] or '*** NOT DETECTED — set manually ***'}")

    print(f"\n Placeholders remaining: {len(rep['placeholders'])}  (NOT auto-removed)")
    for x in dict.fromkeys(rep["placeholders"]):
        print(f"   ! {x}")

    print("\n" + "=" * 62)
    if rep["placeholders"]:
        print(" RESULT: FINALIZED — placeholders remain, human review required")
        print("=" * 62); sys.exit(2)
    print(" RESULT: FINALIZED — no placeholders detected")
    print("=" * 62); sys.exit(0)


if __name__ == "__main__":
    main()
